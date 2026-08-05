#!/usr/bin/env python3
"""Génère le document Word de brief juridique HealthTech — loi n°2013-450 / ARTCI."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Couleurs ──────────────────────────────────────────────────────────────────
ORANGE      = RGBColor(0xE6, 0x6C, 0x00)   # titre principal
DARK_BLUE   = RGBColor(0x1F, 0x35, 0x64)   # en-têtes de section
MID_BLUE    = RGBColor(0x2E, 0x74, 0xB5)   # accent secondaire
GREY_BG     = RGBColor(0xF2, 0xF2, 0xF2)   # fond cellule entête tableau
RED_ECART   = RGBColor(0xC0, 0x00, 0x00)   # écarts critiques
GREEN_OK    = RGBColor(0x37, 0x86, 0x27)   # contrôle en place

def set_cell_bg(cell, rgb_hex: str):
    """Applique une couleur de fond à une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_para(doc, text, bold=False, italic=False, size=10, color=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def table_header_row(table, headers, bg='1F3564'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.alignment = align
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

# ═════════════════════════════════════════════════════════════════════════════
# DONNÉES
# ═════════════════════════════════════════════════════════════════════════════

EXIGENCES = [
    # (REQ, Cat, M/S, Exigence complète, Note pour le juriste, ECART associé)
    ("REQ-LEX-01", "Formalités préalables", "Must",
     "Accomplir la formalité préalable auprès de l'ARTCI (déclaration ou autorisation selon la nature du traitement) AVANT la mise en œuvre.",
     "Confirmer la formalité applicable : déclaration simple ou autorisation ? La nature des données de santé peut imposer une autorisation préalable (voir REQ-LEX-02).",
     None),
    ("REQ-LEX-02", "Formalités / Données sensibles", "Must",
     "Les données de santé étant des données sensibles, le traitement relève d'un régime renforcé (vraisemblablement autorisation préalable ARTCI, et non simple déclaration).",
     "Trancher : autorisation préalable ou déclaration ? — détermine toute la procédure ARTCI. Voir ECART-05.",
     "ECART-05"),
    ("REQ-LEX-03", "Base légale & consentement", "Must",
     "Le traitement doit reposer sur une base légale valable — ici : consentement de la personne concernée.",
     "Confirmer que le consentement est la base légale retenue. Vérifier s'il existe d'autres bases légales mobilisables (obligation légale, intérêt légitime).",
     None),
    ("REQ-LEX-04", "Base légale & consentement", "Must",
     "Le consentement doit être libre, spécifique, éclairé (et univoque), avec preuve de recueil.",
     "Valider le contenu et la forme des écrans de consentement dans l'application (issue #7). Confirmer le format de la preuve de recueil.",
     None),
    ("REQ-LEX-06", "Principes relatifs au traitement", "Must",
     "Finalité déterminée, légitime et explicite ; pas de traitement ultérieur incompatible.",
     "Valider la formulation de la finalité dans le registre des traitements. S'assurer que le partage avec les professionnels de santé via QR code est couvert.",
     None),
    ("REQ-LEX-07", "Principes relatifs au traitement", "Must",
     "Minimisation : données adéquates, pertinentes et limitées à ce qui est nécessaire.",
     "Confirmer que seul le n° CMU, le téléphone et le dossier médical sont collectés. Le modèle zero-knowledge (blob opaque côté serveur) appuie ce principe.",
     None),
    ("REQ-LEX-08", "Principes relatifs au traitement", "Must",
     "Exactitude : données exactes et tenues à jour.",
     "Valider le mécanisme de mise à jour du dossier par le professionnel de santé (issue #18) et par le patient.",
     None),
    ("REQ-LEX-09", "Principes relatifs au traitement", "Must",
     "Loyauté & licéité de la collecte et du traitement.",
     "Revoir les conditions générales d'utilisation et la politique de confidentialité pour en confirmer la conformité.",
     None),
    ("REQ-LEX-10", "Durée de conservation", "Must",
     "Conservation limitée à la durée nécessaire aux finalités, puis suppression/anonymisation — à arbitrer avec les durées minimales de conservation des dossiers médicaux.",
     "Fixer les durées de rétention par catégorie (dossier médical, logs, sauvegardes). Arbitrer avec les obligations de conservation médicale. Voir ECART-01.",
     "ECART-01"),
    ("REQ-LEX-11", "Droits de la personne", "Must",
     "Droit à l'information : informer la personne (finalité, destinataires, droits, etc.) au moment de la collecte.",
     "Valider le contenu de l'écran d'information avant collecte (#7). Vérifier que toutes les mentions obligatoires sont présentes.",
     None),
    ("REQ-LEX-12", "Droits de la personne", "Must",
     "Droit d'accès : la personne peut obtenir communication de ses données.",
     "Le modèle local-first permet au patient d'accéder à ses données directement sur l'appareil. Confirmer si un canal d'accès formel (demande écrite) est également requis.",
     None),
    ("REQ-LEX-13", "Droits de la personne", "Must",
     "Droit de rectification : correction des données inexactes.",
     "Le patient peut-il modifier lui-même ses données ? Valider le périmètre de rectification autorisé au patient vs au professionnel de santé.",
     None),
    ("REQ-LEX-14", "Droits de la personne", "Must",
     "Droit d'opposition (pour motifs légitimes).",
     "Aucun mécanisme d'opposition n'est conçu à ce stade. À concevoir. Voir ECART-02.",
     "ECART-02"),
    ("REQ-LEX-15", "Droits de la personne", "Must",
     "Droit à la suppression / effacement des données.",
     "La solution technique envisagée est le crypto-effacement (destruction de clé = données illisibles). Valider juridiquement si cette approche satisfait l'obligation légale d'effacement. Voir ECART-02.",
     "ECART-02"),
    ("REQ-LEX-16", "Sécurité & confidentialité", "Must",
     "Mesures techniques et organisationnelles assurant confidentialité, intégrité et sécurité des données.",
     "Les mesures techniques sont en place (AES-256-GCM, zero-knowledge, keystore matériel). Valider les mesures organisationnelles (politique de sécurité, gestion des accès internes).",
     None),
    ("REQ-LEX-17", "Sécurité & confidentialité", "Must",
     "Confidentialité / secret des données de santé (secret professionnel).",
     "Confirmer l'articulation avec le secret médical du professionnel de santé. La solution garantit que le serveur ne voit jamais les données en clair.",
     None),
    ("REQ-LEX-18", "Sécurité & confidentialité", "Must",
     "Contrôle des accès aux données ; accès limité aux personnes autorisées (ici : contrôlé par le patient).",
     "Le patient contrôle l'accès via QR éphémère (120 secondes). Confirmer que ce modèle satisfait l'exigence légale de contrôle des accès.",
     None),
    ("REQ-LEX-19", "Résidence & transferts", "Must",
     "Résidence des données sur le territoire national (hébergement en Côte d'Ivoire).",
     "Confirmer la base légale exacte de l'obligation de localisation (obligation statutaire vs mitigation des restrictions de transfert). Voir ECART-07. Nécessaire pour finaliser l'attestation de localisation.",
     "ECART-07"),
    ("REQ-LEX-20", "Résidence & transferts", "Must",
     "Encadrement des transferts transfrontaliers : pas de transfert hors du pays sans garanties adéquates / autorisation.",
     "Confirmer que les dépendances techniques (services tiers, CDN, analytics) ne constituent pas des transferts non autorisés. Voir ECART-07.",
     "ECART-07"),
    ("REQ-LEX-21", "Accountability", "Must",
     "Tenir un registre des activités de traitement (accountability).",
     "Un registre des traitements est rédigé (projet). Le valider et le compléter. Confirmer si la désignation d'un DPO est requise. Voir ECART-06.",
     "ECART-06"),
    ("REQ-LEX-23", "Violations de données", "Must [à confirmer]",
     "Notification des violations de données (à l'ARTCI et/ou aux personnes concernées) dans les délais prescrits.",
     "Confirmer si la loi impose une obligation de notification, à qui, dans quels délais. Concevoir la procédure d'incident. Voir ECART-03.",
     "ECART-03"),
    ("REQ-LEX-24", "Sous-traitance", "Must",
     "Encadrement contractuel de la sous-traitance (hébergeur) : obligations de sécurité et de confidentialité opposables.",
     "Dès la sélection de l'hébergeur souverain (issue #8), prévoir des clauses contractuelles de protection des données. Modèle de clauses à préparer.",
     None),
    ("REQ-LEX-25", "Sécurité (journalisation)", "Must",
     "Aucune donnée médicale en clair, clé de chiffrement, ou PII dans les journaux applicatifs et d'infrastructure.",
     "Valider la politique de journalisation et confirmer que les logs sont couverts par les engagements de confidentialité.",
     None),
]

ECARTS = [
    ("ECART-01", "REQ-LEX-10", "Gouvernance / Juridique",
     "Pas de politique de rétention définie. Tension entre droit à l'oubli et durée minimale légale de conservation des dossiers médicaux non arbitrée.",
     "Définir les durées de conservation par catégorie de données. Fixer le mécanisme de purge. Arbitrer avec les obligations légales de conservation médicale."),
    ("ECART-02", "REQ-LEX-14, REQ-LEX-15", "Conception / Juridique",
     "Aucun mécanisme de suppression / effacement conçu. Acceptabilité juridique du crypto-effacement (destruction de la clé de chiffrement) comme équivalent d'un effacement physique non validée.",
     "Concevoir le flux d'effacement (suppression du blob par UUID + destruction de la clé). Valider juridiquement si le crypto-effacement satisfait l'obligation légale."),
    ("ECART-03", "REQ-LEX-23", "Conception / Gouvernance",
     "Aucune procédure de notification de violation de données en place (délais, destinataires ARTCI / personnes concernées).",
     "Rédiger le runbook de gestion d'incident. Préparer le modèle de notification ARTCI. Confirmer les délais légaux applicables."),
    ("ECART-04", "REQ-LEX-22", "Gouvernance / Juridique",
     "La désignation d'un correspondant / DPO n'est pas encore décidée (requise ou recommandée par l'ARTCI ?).",
     "Confirmer si la désignation est obligatoire ou facultative. Si obligatoire, désigner et consigner dans un acte de gouvernance."),
    ("ECART-05", "REQ-LEX-02, REQ-LEX-05", "Juridique",
     "Régime applicable aux données de santé (autorisation préalable ARTCI vs déclaration simple) non tranché. Régime de consentement des mineurs non confirmé.",
     "Faire trancher par le conseil juridique. La réponse conditionne toute la procédure ARTCI (#30) et les écrans de consentement (#7)."),
    ("ECART-06", "REQ-LEX-21, REQ-LEX-24", "Juridique / Gouvernance",
     "La répartition des rôles responsable de traitement / sous-traitant entre le patient, le professionnel de santé, la plateforme et l'hébergeur n'est pas déterminée.",
     "Qualifier les rôles juridiques de chaque acteur. Mettre à jour le registre des traitements et les modèles de contrats en conséquence."),
    ("ECART-07", "REQ-LEX-19, REQ-LEX-20", "Juridique",
     "La base légale exacte de l'obligation de localisation stricte en Côte d'Ivoire n'est pas précisée (obligation statutaire dure vs mitigation des restrictions de transfert transfrontalier).",
     "Faire trancher par le conseil juridique. La réponse impacte la formulation des exigences de résidence et le libellé de l'attestation de localisation des données."),
    ("ECART-08", "Hors périmètre direct", "Conception / Juridique",
     "Accès d'urgence (break-glass) : patient inconscient ou incapable vs accès strictement contrôlé par QR éphémère. Potentiellement exigé par la régulation santé.",
     "Analyser l'admissibilité juridique d'un accès d'urgence. Toute solution ne doit pas introduire de porte dérobée côté serveur (invariant zero-knowledge)."),
]

SHOULD_EXIGENCES = [
    ("REQ-LEX-05", "Base légale & consentement", "Should [à confirmer]",
     "Mineurs / personnes protégées : régime de consentement spécifique (représentant légal).",
     "Confirmer si la plateforme est ouverte aux mineurs. Si oui, définir le régime de consentement applicable. Voir ECART-05."),
    ("REQ-LEX-22", "Accountability", "Should [à confirmer]",
     "Désignation d'un correspondant / DPO (si requis ou recommandé par l'ARTCI).",
     "Confirmer si la désignation est obligatoire ou recommandée. Voir ECART-04."),
]

# ═════════════════════════════════════════════════════════════════════════════
# GÉNÉRATION DU DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Marges ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style de base ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HEALTHTECH')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Brief de conformité juridique')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Loi n°2013-450 du 19 juin 2013 relative à la protection des données à caractère personnel')
run.font.size = Pt(12)
run.font.color.rgb = MID_BLUE
run.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Exigences Must — Validation juridique requise')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
# fond de paragraphe non supporté nativement — on passe par un tableau 1×1
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tbl.rows[0].cells[0]
set_cell_bg(cell, '1F3564')
cell.width = Cm(14)
p2 = cell.paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Exigences Must — Validation juridique requise')
run2.bold = True
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph()

# Métadonnées
meta_tbl = doc.add_table(rows=5, cols=2)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
meta_data = [
    ('Projet',       'HealthTech — Plateforme de santé numérique décentralisée, Côte d\'Ivoire'),
    ('Date',         '7 juillet 2026'),
    ('Émetteur',     'Équipe technique HealthTech'),
    ('Destinataires','Conseil juridique / Équipe conformité'),
    ('Version',      '1.0 — À valider'),
]
for i, (k, v) in enumerate(meta_data):
    row = meta_tbl.rows[i]
    cell_text(row.cells[0], k, bold=True, size=9, color=DARK_BLUE)
    cell_text(row.cells[1], v, size=9)
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — CONTEXTE ET OBJECTIF
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Contexte et objectif du document', level=1, color=DARK_BLUE)

add_para(doc,
    'HealthTech est une plateforme de santé numérique décentralisée développée pour la Côte d\'Ivoire. '
    'Elle permet au patient de porter son dossier médical chiffré sur son smartphone et d\'en contrôler '
    'l\'accès via un QR code éphémère. L\'architecture est de type « local-first / zero-knowledge » : '
    'les données sont chiffrées côté patient (AES-256-GCM) avant tout transit réseau, et le serveur '
    'ne stocke jamais de données en clair ni de clés de chiffrement.',
    size=10)

doc.add_paragraph()
add_para(doc,
    'Ce document a été préparé par l\'équipe technique à l\'intention de l\'équipe juridique. '
    'Il présente les 22 exigences Must identifiées au titre de la loi n°2013-450 du 19 juin 2013 '
    'et des exigences de l\'ARTCI, et liste les points précis qui nécessitent une validation ou '
    'une décision juridique.',
    size=10)

doc.add_paragraph()

# Encadré objectif
obj_tbl = doc.add_table(rows=1, cols=1)
cell = obj_tbl.rows[0].cells[0]
set_cell_bg(cell, 'E6F0FF')
p_obj = cell.paragraphs[0]
run_obj = p_obj.add_run('Objectif de la revue juridique')
run_obj.bold = True
run_obj.font.size = Pt(10)
run_obj.font.color.rgb = DARK_BLUE
cell.add_paragraph()
p2 = cell.add_paragraph()
run2 = p2.add_run(
    'Valider chacune des 22 exigences Must en apposant un verdict (Validé / Validé avec réserves / Refusé), '
    'une date et le nom du réviseur dans le journal de validation juridique. '
    'La matrice de conformité n\'est considérée comme validée qu\'une fois toutes les exigences Must '
    'signées sans réserve bloquante — condition sine qua non pour la soumission du dossier ARTCI (#30).'
)
run2.font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — ARCHITECTURE DE LA SOLUTION (résumé)
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Architecture de la solution (résumé)', level=1, color=DARK_BLUE)

add_para(doc, 'Comprendre l\'architecture est nécessaire pour évaluer certaines exigences (contrôle des accès, crypto-effacement, résidence des données).', size=10)
doc.add_paragraph()

arch_items = [
    ('Patient', 'L\'application mobile (Android) stocke le dossier médical chiffré localement. La clé maîtresse est générée sur l\'appareil et scellée dans le keystore matériel (TEE/StrongBox). Elle ne quitte jamais l\'appareil en clair.'),
    ('Accès médecin', 'Le patient génère un QR code contenant une clé de session symétrique et l\'URL du blob chiffré sur le serveur. Ce QR expire après 120 secondes. Le médecin déchiffre le dossier en mémoire vive uniquement, sans écriture sur disque.'),
    ('Serveur', 'Le serveur ne stocke que des blobs chiffrés opaques indexés par des UUID anonymes. Il ne dispose d\'aucune clé de déchiffrement et ne peut pas lire les données médicales.'),
    ('Hébergement', 'Le serveur doit être hébergé sur le territoire ivoirien (exigence de résidence — REQ-LEX-19). L\'opérateur souverain est en cours de sélection (issue #8).'),
]

for actor, desc in arch_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run_a = p.add_run(f'{actor} : ')
    run_a.bold = True
    run_a.font.size = Pt(10)
    run_a.font.color.rgb = MID_BLUE
    run_b = p.add_run(desc)
    run_b.font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — TABLEAU DE BORD DES STATUTS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. Tableau de bord — état de la validation', level=1, color=DARK_BLUE)

add_para(doc, 'Au 7 juillet 2026, aucune exigence Must n\'a encore fait l\'objet d\'un sign-off juridique.', size=10)
doc.add_paragraph()

dash_tbl = doc.add_table(rows=6, cols=3)
dash_tbl.style = 'Table Grid'
headers = ['Indicateur', 'Valeur', 'Observation']
table_header_row(dash_tbl, headers)
rows_data = [
    ('Exigences Must totales', '22', 'REQ-LEX-01 à 25 (hors REQ-LEX-05 et REQ-LEX-22 classés Should)'),
    ('Exigences Should / à confirmer', '3', 'REQ-LEX-05, REQ-LEX-22, REQ-LEX-23'),
    ('Exigences Must validées', '0 / 22', 'Aucune — revue en attente'),
    ('Écarts ouverts', '8', 'ECART-01 à ECART-08 — détail section 5'),
    ('Statut de la matrice', 'PROJET', 'Non validée — critère #5 non atteint'),
]
for i, (ind, val, obs) in enumerate(rows_data):
    row = dash_tbl.rows[i + 1]
    cell_text(row.cells[0], ind, bold=True, size=9)
    is_zero = '0' in val or 'PROJET' in val
    cell_text(row.cells[1], val, bold=True, size=9,
              color=RED_ECART if is_zero else GREEN_OK)
    cell_text(row.cells[2], obs, size=9)
    row.cells[0].width = Cm(5.5)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(8)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — LES 22 EXIGENCES MUST
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Les 22 exigences Must — detail et questions pour le conseil juridique', level=1, color=DARK_BLUE)

add_para(doc,
    'Pour chaque exigence, le conseil juridique est invité à (1) confirmer ou corriger la formulation, '
    '(2) préciser le numéro d\'article exact de la loi n°2013-450, et (3) apposer son verdict dans le '
    'journal de validation (fichier : docs/compliance/journal-validation-juridique.md).',
    size=10, italic=True)
doc.add_paragraph()

# Groupes
groups = {
    'A. Formalités préalables': ['REQ-LEX-01', 'REQ-LEX-02'],
    'B. Base légale & consentement': ['REQ-LEX-03', 'REQ-LEX-04'],
    'C. Principes relatifs au traitement': ['REQ-LEX-06', 'REQ-LEX-07', 'REQ-LEX-08', 'REQ-LEX-09'],
    'D. Durée de conservation': ['REQ-LEX-10'],
    'E. Droits de la personne concernée': ['REQ-LEX-11', 'REQ-LEX-12', 'REQ-LEX-13', 'REQ-LEX-14', 'REQ-LEX-15'],
    'F. Sécurité & confidentialité': ['REQ-LEX-16', 'REQ-LEX-17', 'REQ-LEX-18'],
    'G. Résidence & transferts transfrontaliers': ['REQ-LEX-19', 'REQ-LEX-20'],
    'H. Accountability': ['REQ-LEX-21'],
    'I. Violations de données': ['REQ-LEX-23'],
    'J. Sous-traitance': ['REQ-LEX-24'],
    'K. Journalisation': ['REQ-LEX-25'],
}

exig_map = {e[0]: e for e in EXIGENCES}

for group_title, req_ids in groups.items():
    add_heading(doc, f'  {group_title}', level=2, color=MID_BLUE)

    for req_id in req_ids:
        if req_id not in exig_map:
            continue
        req, cat, ms, exigence, note, ecart = exig_map[req_id]

        # Titre de l'exigence
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(f'{req} — {cat}')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_BLUE

        # Badge Must
        run_ms = p.add_run(f'   [{ms}]')
        run_ms.bold = True
        run_ms.font.size = Pt(9)
        run_ms.font.color.rgb = RED_ECART if 'confirmer' in ms.lower() else RGBColor(0xC0, 0x00, 0x00)

        # Texte de l'exigence
        req_tbl = doc.add_table(rows=2, cols=1)
        req_tbl.style = 'Table Grid'

        c1 = req_tbl.rows[0].cells[0]
        set_cell_bg(c1, 'F2F2F2')
        p1 = c1.paragraphs[0]
        r1a = p1.add_run('Exigence : ')
        r1a.bold = True
        r1a.font.size = Pt(9)
        r1b = p1.add_run(exigence)
        r1b.font.size = Pt(9)

        c2 = req_tbl.rows[1].cells[0]
        set_cell_bg(c2, 'FFF8DC' if ecart else 'EBF3E8')
        p2 = c2.paragraphs[0]
        r2a = p2.add_run('Question / Action pour le conseil juridique : ')
        r2a.bold = True
        r2a.font.size = Pt(9)
        r2a.font.color.rgb = DARK_BLUE
        r2b = p2.add_run(note)
        r2b.font.size = Pt(9)
        if ecart:
            p2.add_run(f'  → {ecart}').font.color.rgb = RED_ECART

        doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4B — EXIGENCES SHOULD
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4b. Exigences Should / À confirmer', level=1, color=DARK_BLUE)
add_para(doc, 'Ces exigences ne bloquent pas la validation de la matrice Must, mais requièrent une décision ou une confirmation juridique.', size=10, italic=True)
doc.add_paragraph()

for req, cat, ms, exigence, note in SHOULD_EXIGENCES:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f'{req} — {cat}')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MID_BLUE
    p.add_run(f'   [{ms}]').font.color.rgb = RGBColor(0x80, 0x40, 0x00)

    t = doc.add_table(rows=2, cols=1)
    t.style = 'Table Grid'
    c1 = t.rows[0].cells[0]
    set_cell_bg(c1, 'F2F2F2')
    p1 = c1.paragraphs[0]
    p1.add_run('Exigence : ').bold = True
    p1.runs[0].font.size = Pt(9)
    p1.add_run(exigence).font.size = Pt(9)
    c2 = t.rows[1].cells[0]
    set_cell_bg(c2, 'FFF8DC')
    p2 = c2.paragraphs[0]
    r = p2.add_run('Question / Action : ')
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = DARK_BLUE
    p2.add_run(note).font.size = Pt(9)
    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — ÉCARTS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Registre des écarts — points critiques à instruire', level=1, color=DARK_BLUE)

add_para(doc,
    'Les écarts ci-dessous représentent des obligations légales sans contrôle technique conçu ou '
    'des questions juridiques non tranchées. Chaque écart bloque la conformité des exigences Must '
    'associées. La résolution passe obligatoirement par une décision ou un avis juridique.',
    size=10)
doc.add_paragraph()

ecart_tbl = doc.add_table(rows=len(ECARTS) + 1, cols=5)
ecart_tbl.style = 'Table Grid'
table_header_row(ecart_tbl, ['ID', 'Exigences impactées', 'Type', 'Écart constaté', 'Action proposée'])
set_col_widths(ecart_tbl, [1.5, 2.5, 2.2, 5.5, 5.5])

for i, (eid, reqs, typ, ecart_desc, action) in enumerate(ECARTS):
    row = ecart_tbl.rows[i + 1]
    is_crit = eid in ['ECART-01', 'ECART-02', 'ECART-03', 'ECART-05', 'ECART-07']
    cell_text(row.cells[0], eid, bold=True, size=8, color=RED_ECART if is_crit else DARK_BLUE)
    cell_text(row.cells[1], reqs, size=8)
    cell_text(row.cells[2], typ, size=8)
    cell_text(row.cells[3], ecart_desc, size=8)
    cell_text(row.cells[4], action, size=8)
    if is_crit:
        set_cell_bg(row.cells[0], 'FFE0E0')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — JOURNAL DE VALIDATION (template)
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Journal de validation juridique — à compléter', level=1, color=DARK_BLUE)

add_para(doc,
    'Ce tableau doit être complété par le conseil juridique. Pour chaque exigence Must : '
    'apposer le verdict, la date et le nom du réviseur. Une réserve bloquante empêche '
    'la clôture de la matrice.',
    size=10, italic=True)
doc.add_paragraph()

all_reqs_for_journal = [(r, ms) for r, _, ms, _, _, _ in EXIGENCES]
journal_tbl = doc.add_table(rows=len(all_reqs_for_journal) + 1, cols=5)
journal_tbl.style = 'Table Grid'
table_header_row(journal_tbl, ['REQ', 'Must/Should', 'Verdict', 'Date', 'Réviseur / Commentaires'])
set_col_widths(journal_tbl, [2.0, 2.0, 2.5, 2.0, 8.5])

for i, (req, ms) in enumerate(all_reqs_for_journal):
    row = journal_tbl.rows[i + 1]
    cell_text(row.cells[0], req, bold=True, size=9)
    cell_text(row.cells[1], ms, size=9)
    cell_text(row.cells[2], 'En attente', size=9, color=RGBColor(0x80, 0x80, 0x80))
    cell_text(row.cells[3], '', size=9)
    cell_text(row.cells[4], '', size=9)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — PROCHAINES ÉTAPES
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. Prochaines étapes recommandées', level=1, color=DARK_BLUE)

steps = [
    ('Priorité 1 — Formalité ARTCI (REQ-LEX-01/02)',
     'Confirmer le régime applicable (déclaration vs autorisation préalable pour données de santé). '
     'C\'est le premier prérequis de toute la procédure d\'homologation. Instruire ECART-05.'),
    ('Priorité 2 — Écarts critiques (ECART-01, 02, 03)',
     'Définir la politique de rétention (ECART-01), concevoir le flux d\'effacement et valider '
     'le crypto-effacement (ECART-02), rédiger la procédure de notification de violation (ECART-03).'),
    ('Priorité 3 — Qualification des rôles RT / sous-traitant (ECART-06)',
     'Déterminer les rôles juridiques des acteurs (patient, médecin, plateforme, hébergeur). '
     'Conditionne le registre des traitements et les clauses contractuelles.'),
    ('Priorité 4 — Base légale de la résidence (ECART-07)',
     'Préciser la base légale de l\'obligation de localisation en Côte d\'Ivoire. '
     'Nécessaire pour finaliser l\'attestation de localisation et choisir l\'hébergeur (#8).'),
    ('Priorité 5 — Revue exigence par exigence',
     'Parcourir le journal (section 6) et apposer le verdict sur chacune des 22 exigences Must. '
     'Confirmer les numéros d\'articles exact de la loi n°2013-450.'),
    ('Priorité 6 — Désignation DPO / correspondant (ECART-04)',
     'Décider si la désignation est obligatoire ou facultative, et si oui, procéder à la désignation.'),
]

for titre, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run_t = p.add_run(f'▶  {titre}')
    run_t.bold = True
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = DARK_BLUE
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.0)
    run_d = p2.add_run(desc)
    run_d.font.size = Pt(10)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# PIED DE PAGE / NOTE FINALE
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

add_heading(doc, 'Note de l\'équipe technique', level=1, color=DARK_BLUE)

note_tbl = doc.add_table(rows=1, cols=1)
cell = note_tbl.rows[0].cells[0]
set_cell_bg(cell, 'E6F0FF')
p_note = cell.paragraphs[0]
run_note = p_note.add_run(
    'Invariant zero-knowledge — engagement de l\'équipe technique\n\n'
    'Quelle que soit la résolution des écarts juridiques, l\'architecture HealthTech maintient '
    'l\'invariant suivant : le serveur ne déchiffre jamais les données médicales et ne détient '
    'aucune clé. Toute demande de contrôle légal ou d\'accès aux données passe par le patient '
    '(porteur de la clé maîtresse). Si une exigence juridique semblait imposer une porte dérobée '
    'côté serveur, l\'équipe technique souhaite en être informée pour évaluer l\'impact sur '
    'l\'architecture avant toute décision.'
)
run_note.font.size = Pt(10)
run_note.font.color.rgb = DARK_BLUE

doc.add_paragraph()
add_para(doc,
    'Documents de référence disponibles dans le dépôt GitHub (kortiene/HealthTech) :',
    bold=True, size=10)
refs = [
    'docs/compliance/exigences-legales.md — registre des exigences REQ-LEX-NN',
    'docs/compliance/loi-2013-450-artci-matrix.md — matrice de conformité complète',
    'docs/compliance/journal-validation-juridique.md — journal de sign-off (à compléter)',
    'docs/compliance/ecarts.md — registre des écarts ECART-NN',
    'docs/compliance/registre-des-traitements.md — registre des traitements',
    'docs/compliance/attestation-localisation-donnees.md — modèle d\'attestation',
    'docs/compliance/homologation-artci/ — dossier d\'homologation consolidé',
]
for ref in refs:
    add_bullet(doc, ref, size=9)

doc.add_paragraph()
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_final = p_final.add_run('HealthTech — Document confidentiel — usage interne — juillet 2026')
run_final.font.size = Pt(8)
run_final.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run_final.italic = True

# ── Sauvegarde ────────────────────────────────────────────────────────────────
output_path = '/Users/macbook/Documents/app/Health/HealthTech/docs/compliance/HealthTech_Brief_Juridique_ARTCI_2026-07.docx'
doc.save(output_path)
print(f'Document généré : {output_path}')
